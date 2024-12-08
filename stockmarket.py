from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
import os
import time
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import re
import glob
import requests
import base64
from dotenv import load_dotenv # type: ignore
from groq import Client # type: ignore

# Load environment variables
load_dotenv()

# Configure Groq client
client = Client(api_key=os.getenv('GROQ_API_KEY'))

# Path to your chromedriver
PATH = "C:/Users/ajays/Documents/chromedriver-win64/chromedriver.exe"

# Define the blog directory path
BLOG_DIR = r"C:\Users\ajays\Desktop\CapX\twitter content generation"

def ensure_blog_directory():
    """Create the blog directory if it doesn't exist"""
    if not os.path.exists(BLOG_DIR):
        os.makedirs(BLOG_DIR)
        print(f"Created blog directory at: {BLOG_DIR}")

def login_to_twitter(driver, username, password):
    """Log in to Twitter using provided credentials"""
    driver.get("https://twitter.com/login")
    username_field = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, "//input[@name='text']"))
    )
    username_field.send_keys(username)

    next_button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//span[contains(text(),'Next')]"))
    )
    next_button.click()

    time.sleep(15)

    password_field = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, "//input[@name='password']"))
    )
    password_field.send_keys(password)

    log_in = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.XPATH, "//span[contains(text(),'Log in')]"))
    )
    log_in.click()

    WebDriverWait(driver, 15).until(EC.url_contains("home"))

def search_latest_ai_news(driver, search_url):
    """Search for AI-related tweets"""
    driver.get(search_url)
    time.sleep(5)

def scroll_and_load_tweets(driver):
    """Scroll to load more tweets"""
    try:
        driver.find_element(By.TAG_NAME, "body").send_keys(Keys.PAGE_DOWN)
        time.sleep(2)
    except Exception as e:
        print(f"Error during scroll: {e}")

def is_tweet_relevant(tweet, keywords):
    """Check if tweet contains relevant keywords"""
    return any(keyword.lower() in tweet.lower() for keyword in keywords)

def clean_tweet_text(text):
    """Clean tweet text by removing URLs, mentions, hashtags, etc."""
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'@\w+', '', text)
    text = re.sub(r'#\w+', '', text)
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    text = ' '.join(text.split())
    return text

def extract_new_tweets(driver, used_tweets, keywords):
    """Extract new relevant tweets, up to 5 per keyword"""
    tweets_by_keyword = {}
    
    for keyword in keywords:
        tweets_by_keyword[keyword] = []
        max_attempts = 10  # More attempts to find multiple tweets
        
        while len(tweets_by_keyword[keyword]) < 5 and max_attempts > 0:
            tweet_elements = driver.find_elements(By.XPATH, "//article[@data-testid='tweet']")
            
            for tweet in tweet_elements:
                try:
                    tweet_text = tweet.find_element(By.XPATH, ".//div[@data-testid='tweetText']").text
                    if tweet_text not in used_tweets and is_tweet_relevant(tweet_text, [keyword]):
                        clean_text = clean_tweet_text(tweet_text)
                        if clean_text not in tweets_by_keyword[keyword]:
                            tweets_by_keyword[keyword].append(clean_text)
                            if len(tweets_by_keyword[keyword]) >= 5:
                                break
                except Exception as e:
                    print(f"Error extracting tweet: {e}")
                    continue
            
            if len(tweets_by_keyword[keyword]) < 5:
                scroll_and_load_tweets(driver)
                max_attempts -= 1
    
    return tweets_by_keyword

def generate_image_prompts(blog_content, keyword):
    """Generate better image prompts for different sections of the blog"""
    sections = blog_content.split('\n\n')
    image_prompts = []
    
    # Get the title from the first line
    title = sections[0].strip().replace('#', '').strip()
    
    # First image - realistic future vision
    intro_prompt = (
        f"Photorealistic image of {keyword} in a modern setting. "
        f"Style: High-quality photography, 4K, detailed. "
        f"Include: Real environment and natural lighting. "
        f"Theme: Professional and sophisticated. "
        f"Mood: Premium and polished. "
        f"Photography style: Commercial product photography."
    )
    image_prompts.append(("intro", intro_prompt))
    
    # Second image - real-world usage
    middle_prompt = (
        f"Documentary-style photograph of {keyword} in use. "
        f"Style: Candid photography, natural lighting. "
        f"Include: Real people interacting with {keyword}. "
        f"Theme: Everyday life and practical applications. "
        f"Mood: Authentic and relatable. "
        f"Photography style: Photojournalistic."
    )
    image_prompts.append(("middle", middle_prompt))
    
    # Closing image - realistic future impact
    conclusion_prompt = (
        f"Professional photograph showcasing {keyword} in action. "
        f"Style: High-end editorial photography. "
        f"Include: Dynamic composition and dramatic real lighting. "
        f"Theme: Innovation in real settings. "
        f"Mood: Impactful and inspiring. "
        f"Photography style: Magazine cover quality."
    )
    image_prompts.append(("conclusion", conclusion_prompt))
    
    return image_prompts

class HuggingFaceImageGenerator:
    def __init__(self):
        try:
            self.api_token = os.getenv('HF_API_TOKEN')
            if not self.api_token:
                raise ValueError("HF_API_TOKEN not found in environment variables")
            
            self.headers = {
                "Authorization": f"Bearer {self.api_token}",
                "Content-Type": "application/json"
            }
            # Using a different model that's typically more available
            self.api_url = "https://api-inference.huggingface.co/models/stabilityai/stable-diffusion-2-1"
            self.image_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_images")
            os.makedirs(self.image_dir, exist_ok=True)
            
            print("Successfully initialized Hugging Face image generator")
            self.initialized = True
                
        except Exception as e:
            print(f"Error initializing Hugging Face generator: {e}")
            self.initialized = False

    def generate_image(self, prompt):
        try:
            if not self.initialized:
                return None

            print(f"Generating image for prompt: {prompt}")
            
            # Prepare the payload
            payload = {
                "inputs": prompt,
                "parameters": {
                    "negative_prompt": "blurry, bad quality, distorted, ugly, bad art, poor details",
                    "num_inference_steps": 30,
                    "guidance_scale": 7.5,
                    "width": 768,
                    "height": 768
                }
            }
            
            # Make API request with retry logic
            max_retries = 3
            retry_delay = 20  # seconds
            
            for attempt in range(max_retries):
                response = requests.post(self.api_url, headers=self.headers, json=payload)
                
                if response.status_code == 200:
                    # Save the image
                    timestamp = int(time.time())
                    image_path = os.path.join(self.image_dir, f"generated_image_{timestamp}.png")
                    
                    with open(image_path, 'wb') as f:
                        f.write(response.content)
                    print(f"Image saved to: {image_path}")
                    return image_path
                else:
                    error_data = response.json()
                    if "error" in error_data and "estimated_time" in error_data:
                        wait_time = min(retry_delay, int(error_data["estimated_time"]) + 5)
                        print(f"Model loading... Waiting {wait_time} seconds before retry. Attempt {attempt + 1}/{max_retries}")
                        time.sleep(wait_time)
                    else:
                        print(f"Error generating image: {response.text}")
                        if attempt < max_retries - 1:
                            print(f"Retrying in {retry_delay} seconds... Attempt {attempt + 1}/{max_retries}")
                            time.sleep(retry_delay)
                        else:
                            return None
            
            print("Failed to generate image after all retries")
            return None
            
        except Exception as e:
            print(f"Error generating image: {e}")
            return None

def generate_blogs(tweets_by_keyword, max_retries=3, retry_delay=5):
    
    """Generate comprehensive blog posts using Groq API with Mixtral model"""
    blogs = []
    image_generator = HuggingFaceImageGenerator()  # Using Hugging Face
    
    for keyword, tweets in tweets_by_keyword.items():
        for attempt in range(max_retries):
            try:
                # Create a more focused prompt based on the tweets
                tweet_insights = "\n".join([f"- {tweet}" for tweet in tweets])
                
                # Enhanced prompt with better structure and examples
                prompt = f"""Create an in-depth, professional blog post about {keyword} based on these insights:

{tweet_insights}

Structure your response following this exact format:

# [Write an SEO-optimized, attention-grabbing title]

## Introduction
[Write a compelling 150-word introduction that:
- Starts with a powerful hook or surprising statistic
- Introduces {keyword} and its significance
- Outlines the key themes from the collected insights
- Sets up the main points to be discussed]

## Current State of {keyword}
[Write 250 words analyzing:
- Latest developments and breakthroughs
- Key players and their contributions
- Current challenges and opportunities
- Market trends and adoption rates]

## Technical Deep Dive
[Write 300 words covering:
- Core technologies and methodologies
- Technical specifications and requirements
- Implementation considerations
- Best practices and standards]

## Impact Analysis
[Write 200 words examining:
- Industry implications
- Business transformations
- Economic effects
- Societal changes]

## Future Outlook
[Write 200 words discussing:
- Predicted developments
- Upcoming challenges
- Potential breakthroughs
- Industry predictions]

## Conclusion
[Write a 150-word conclusion that:
- Summarizes key points
- Provides actionable insights
- Ends with a thought-provoking statement
- Encourages further exploration]

Style Requirements:
1. Use clear, authoritative language
2. Include specific examples and data points
3. Break down complex concepts
4. Use industry-standard terminology
5. Maintain a professional tone
6. Format with proper Markdown headings
7. Include relevant statistics when possible
8. Use short paragraphs for readability

Focus on providing actionable insights and practical implications while maintaining technical accuracy."""

                completion = client.chat.completions.create(
                    model="mixtral-8x7b-32768",
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are an expert technology analyst and writer specializing in AI and emerging technologies. "
                                "You excel at creating comprehensive, well-researched blog posts that blend technical depth with "
                                "accessibility. Your writing is known for being data-driven, insightful, and engaging while maintaining "
                                "professional standards and technical accuracy."
                            )
                        },
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=4096
                )

                blog_content = completion.choices[0].message.content
                
                # Generate and add images with better error handling
                image_prompts = generate_image_prompts(blog_content, keyword)
                images = []
                for section, img_prompt in image_prompts:
                    try:
                        image_path = image_generator.generate_image(img_prompt)
                        if image_path and os.path.exists(image_path):
                            images.append((section, image_path))
                        else:
                            print(f"Warning: Failed to generate image for {section}")
                    except Exception as e:
                        print(f"Error generating image for {section}: {e}")
                        continue
                
                blogs.append({"content": blog_content, "images": images, "keyword": keyword})
                break
                
            except Exception as e:
                print(f"Error generating blog (attempt {attempt + 1}/{max_retries}): {e}")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                else:
                    print(f"Failed to generate blog for keyword: {keyword}")
    
    return blogs

def generate_blog_with_images(keyword, tweets):
    """Generate a blog post with images for the given keyword and tweets"""
    try:
        # Initialize Hugging Face image generator
        image_generator = HuggingFaceImageGenerator()
        
        if not image_generator.initialized:
            print("Warning: Image generator not initialized. Proceeding without images.")
            return generate_blogs(keyword, tweets)
        
        # Generate blog content first
        blog_content = generate_blog_content(keyword, tweets)
        if not blog_content:
            return None
        
        # Generate image prompts
        image_prompts = generate_image_prompts(blog_content, keyword)
        images = []
        
        # Generate images
        for section, prompt in image_prompts:
            image_path = image_generator.generate_image(prompt)
            if image_path:
                images.append((section, image_path))
            else:
                print(f"Warning: Failed to generate image for {section}")
        
        return {"keyword": keyword, "content": blog_content, "images": images}
        
    except Exception as e:
        print(f"Error generating blog with images: {e}")
        return None

def generate_blog_content(keyword, tweets):
    """Generate blog content using Groq API with Mixtral model"""
    try:
        # Create a more focused prompt based on the tweets
        tweet_insights = "\n".join([f"- {tweet}" for tweet in tweets])
        
        # Enhanced prompt with better structure and examples
        prompt = f"""Create an in-depth, professional blog post about {keyword} based on these insights:

{tweet_insights}

Structure your response following this exact format:

# [Write an SEO-optimized, attention-grabbing title]

## Introduction
[Write a compelling 150-word introduction that:
- Starts with a powerful hook or surprising statistic
- Introduces {keyword} and its significance
- Outlines the key themes from the collected insights
- Sets up the main points to be discussed]

## Current State of {keyword}
[Write 250 words analyzing:
- Latest developments and breakthroughs
- Key players and their contributions
- Current challenges and opportunities
- Market trends and adoption rates]

## Technical Deep Dive
[Write 300 words covering:
- Core technologies and methodologies
- Technical specifications and requirements
- Implementation considerations
- Best practices and standards]

## Impact Analysis
[Write 200 words examining:
- Industry implications
- Business transformations
- Economic effects
- Societal changes]

## Future Outlook
[Write 200 words discussing:
- Predicted developments
- Upcoming challenges
- Potential breakthroughs
- Industry predictions]

## Conclusion
[Write a 150-word conclusion that:
- Summarizes key points
- Provides actionable insights
- Ends with a thought-provoking statement
- Encourages further exploration]

Style Requirements:
1. Use clear, authoritative language
2. Include specific examples and data points
3. Break down complex concepts
4. Use industry-standard terminology
5. Maintain a professional tone
6. Format with proper Markdown headings
7. Include relevant statistics when possible
8. Use short paragraphs for readability

Focus on providing actionable insights and practical implications while maintaining technical accuracy."""

        completion = client.chat.completions.create(
            model="mixtral-8x7b-32768",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert technology analyst and writer specializing in AI and emerging technologies. "
                        "You excel at creating comprehensive, well-researched blog posts that blend technical depth with "
                        "accessibility. Your writing is known for being data-driven, insightful, and engaging while maintaining "
                        "professional standards and technical accuracy."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4096
        )

        blog_content = completion.choices[0].message.content
        return blog_content
        
    except Exception as e:
        print(f"Error generating blog content: {e}")
        return None

def save_blogs_to_word(blogs):
    """Save generated blogs to Word documents with images"""
    ensure_blog_directory()
    
    for blog in blogs:
        try:
            doc = Document()
            content = blog["content"]
            images = blog["images"]
            keyword = blog["keyword"]
            
            # Track which images have been added
            added_images = {"intro": False, "middle": False, "conclusion": False}
            
            # Split the blog content into sections
            sections = content.split('\n\n')
            
            # Add title
            title = sections[0].strip().replace('#', '').strip()
            doc.add_heading(title, 0)
            
            # Add intro image right after title
            for section, image_path in images:
                if section == "intro" and image_path and os.path.exists(image_path):
                    try:
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(4))
                        print(f"Successfully added intro image from {image_path}")
                        added_images["intro"] = True
                        break
                    except Exception as e:
                        print(f"Error adding intro image: {e}")
            
            # Process each section
            current_section = 1  # Skip title section
            middle_section = len(sections) // 2
            
            while current_section < len(sections):
                section_text = sections[current_section].strip()
                if not section_text:
                    current_section += 1
                    continue
                
                # Check if it's a heading
                if section_text.startswith('#'):
                    level = min(section_text.count('#'), 9)
                    text = section_text.lstrip('#').strip()
                    
                    # Check if this is the conclusion section
                    if "conclusion" in text.lower() and not added_images["conclusion"]:
                        # Add conclusion image before conclusion heading
                        for section, image_path in images:
                            if section == "conclusion" and image_path and os.path.exists(image_path):
                                try:
                                    paragraph = doc.add_paragraph()
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run()
                                    run.add_picture(image_path, width=Inches(4))
                                    print(f"Successfully added conclusion image from {image_path}")
                                    added_images["conclusion"] = True
                                    break
                                except Exception as e:
                                    print(f"Error adding conclusion image: {e}")
                    
                    # Add the heading
                    doc.add_heading(text, level)
                    
                    # Add middle image after the first major heading if not added
                    if current_section >= middle_section and not added_images["middle"]:
                        for section, image_path in images:
                            if section == "middle" and image_path and os.path.exists(image_path):
                                try:
                                    paragraph = doc.add_paragraph()
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run()
                                    run.add_picture(image_path, width=Inches(4))
                                    print(f"Successfully added middle image from {image_path}")
                                    added_images["middle"] = True
                                    break
                                except Exception as e:
                                    print(f"Error adding middle image: {e}")
                else:
                    # Add regular paragraph
                    doc.add_paragraph(section_text)
                
                current_section += 1
            
            # Save the document
            filename = f"blog_{keyword}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(BLOG_DIR, filename)
            doc.save(filepath)
            print(f"Saved blog to: {filepath}")
            
        except Exception as e:
            print(f"Error saving blog: {e}")

def delete_previous_blog_files():
    """Delete previous blog files"""
    ensure_blog_directory()
    for file in glob.glob(os.path.join(BLOG_DIR, "*.docx")):
        try:
            os.remove(file)
            print(f"Deleted: {file}")
        except Exception as e:
            print(f"Error deleting {file}: {e}")

def load_used_tweets(csv_filename):
    """Load previously used tweets from CSV"""
    try:
        df = pd.read_csv(csv_filename)
        # Check if 'tweet' column exists
        if 'tweet' in df.columns:
            return set(df['tweet'].values)
        else:
            print(f"Warning: 'tweet' column not found in {csv_filename}")
            return set()
    except FileNotFoundError:
        print(f"No previous tweets file found at {csv_filename}")
        return set()
    except Exception as e:
        print(f"Error loading tweets: {e}")
        return set()

def save_used_tweet(csv_filename, tweet):
    """Save new tweet to CSV"""
    df = pd.DataFrame({'tweet': [tweet], 'timestamp': [time.strftime('%Y-%m-%d %H:%M:%S')]})
    df.to_csv(csv_filename, mode='a', header=not os.path.exists(csv_filename), index=False)

def get_user_search_terms():
    """Get search terms from user"""
    print("\nEnter search terms (one per line, press Enter twice to finish):")
    terms = []
    while True:
        term = input()
        if term:
            terms.append(term)
        else:
            break
    return terms

def build_search_query(search_terms):
    """Build Twitter search URL from terms"""
    query = ' OR '.join(f'"{term}"' for term in search_terms)
    return f"https://twitter.com/search?q={query}&src=typed_query&f=live"

def main():
    """Main execution flow"""
    ensure_blog_directory()
    
    # Get Twitter credentials from config.py
    twitter_username = os.getenv('TWITTER_USERNAME')
    twitter_password = os.getenv('TWITTER_PASSWORD')
    
    if not all([twitter_username, twitter_password]):
        print("Please set TWITTER_USERNAME and TWITTER_PASSWORD in your .env file")
        return
    
    # Get search terms from user
    search_terms = get_user_search_terms()
    if not search_terms:
        print("No search terms provided. Exiting...")
        return
    
    # Initialize Chrome driver
    service = Service(PATH)
    driver = webdriver.Chrome(service=service)
    
    try:
        # Load previously used tweets
        used_tweets = load_used_tweets('used_tweets.csv')
        
        # Login to Twitter
        login_to_twitter(driver, twitter_username, twitter_password)
        
        # Build and execute search
        search_url = build_search_query(search_terms)
        search_latest_ai_news(driver, search_url)
        
        # Extract tweets and generate blogs
        tweets_by_keyword = extract_new_tweets(driver, used_tweets, search_terms)
        if tweets_by_keyword:
            # Save new tweets to CSV
            for keyword, tweets in tweets_by_keyword.items():
                for tweet in tweets:
                    save_used_tweet('used_tweets.csv', tweet)
            
            # Generate and save blogs
            delete_previous_blog_files()
            blogs = []
            for keyword, tweets in tweets_by_keyword.items():
                blog = generate_blog_with_images(keyword, tweets)
                if blog:
                    blogs.append(blog)
            save_blogs_to_word(blogs)
            
            print("\nBlog generation complete! Check the blogs directory for the new files.")
        else:
            print("\nNo new relevant tweets found. Try again later or modify your search terms.")
    
    finally:
        driver.quit()

if __name__ == "__main__":
    main()